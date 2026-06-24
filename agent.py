import csv
import html
import json
import os
import re
import shlex
import subprocess
from io import StringIO
from html.parser import HTMLParser
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import parse_qs, quote_plus, unquote, urlparse
from urllib.request import Request, urlopen


MAX_FILE_BYTES = 200_000
MAX_TOOL_OUTPUT = 12_000
OUTPUTS_DIR = "outputs"
SEARXNG_URL = os.environ.get(
    "SEARXNG_URL", "http://192.168.1.249:8081"
).rstrip("/")
BLOCKED_COMMANDS = {
    "chmod",
    "chown",
    "curl",
    "dd",
    "docker",
    "git",
    "kill",
    "killall",
    "mount",
    "nc",
    "netcat",
    "npm",
    "pip",
    "pip3",
    "pkill",
    "rm",
    "shutdown",
    "ssh",
    "sudo",
    "wget",
}


class SearchResultParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.results = []
        self.current_url = None
        self.current_text = []

    def handle_starttag(self, tag, attrs):
        attributes = dict(attrs)
        if tag == "a" and "result__a" in attributes.get("class", ""):
            self.current_url = attributes.get("href")
            self.current_text = []

    def handle_data(self, data):
        if self.current_url:
            self.current_text.append(data)

    def handle_endtag(self, tag):
        if tag == "a" and self.current_url:
            self.results.append(
                {
                    "title": " ".join("".join(self.current_text).split()),
                    "url": clean_search_url(self.current_url),
                }
            )
            self.current_url = None
            self.current_text = []


def clean_search_url(url):
    parsed = urlparse(url)
    target = parse_qs(parsed.query).get("uddg")
    return unquote(target[0]) if target else url


def truncate(value, limit=MAX_TOOL_OUTPUT):
    value = str(value)
    if len(value) <= limit:
        return value
    return f"{value[:limit]}\n...[truncated]"


def safe_path(workspace, relative_path):
    root = Path(workspace).resolve()
    candidate = (root / relative_path).resolve()
    if candidate != root and root not in candidate.parents:
        raise ValueError("Path must stay inside the agent workspace.")
    return candidate


def list_files(workspace, path="."):
    target = safe_path(workspace, path)
    if not target.exists():
        return f"Path does not exist: {path}"
    if target.is_file():
        return str(target.relative_to(Path(workspace).resolve()))

    entries = []
    for item in sorted(target.rglob("*")):
        if len(entries) >= 500:
            entries.append("...[file listing truncated]")
            break
        relative = item.relative_to(Path(workspace).resolve())
        entries.append(f"{relative}/" if item.is_dir() else str(relative))
    return "\n".join(entries) or "(empty workspace)"


def read_file(workspace, path):
    target = safe_path(workspace, path)
    if not target.is_file():
        return f"File does not exist: {path}"
    if target.stat().st_size > MAX_FILE_BYTES:
        return f"File is too large to read ({target.stat().st_size} bytes)."
    return truncate(target.read_text(encoding="utf-8", errors="replace"))


def write_file(workspace, path, content):
    target = safe_path(workspace, path)
    target.parent.mkdir(parents=True, exist_ok=True)
    encoded = content.encode("utf-8")
    if len(encoded) > MAX_FILE_BYTES:
        return f"Refused: content exceeds {MAX_FILE_BYTES} bytes."
    target.write_text(content, encoding="utf-8")
    return f"Wrote {len(encoded)} bytes to {target.relative_to(Path(workspace).resolve())}"


def output_path(workspace, filename, extension):
    safe_name = Path(filename).name or f"document{extension}"
    if not safe_name.lower().endswith(extension):
        safe_name = f"{safe_name}{extension}"
    return safe_path(workspace, f"{OUTPUTS_DIR}/{safe_name}")


def create_markdown(workspace, filename, content):
    target = output_path(workspace, filename, ".md")
    target.parent.mkdir(parents=True, exist_ok=True)
    encoded = content.encode("utf-8")
    if len(encoded) > MAX_FILE_BYTES:
        return f"Refused: content exceeds {MAX_FILE_BYTES} bytes."
    target.write_text(content, encoding="utf-8")
    relative = target.relative_to(Path(workspace).resolve())
    return f"Created Markdown document: {relative}\nDownload: /agent_outputs/{target.name}"


def argument_text(arguments):
    return arguments.get("content") or arguments.get("text") or arguments.get("body") or ""


def argument_filename(arguments, default_name, extension):
    filename = arguments.get("filename") or arguments.get("path") or default_name
    filename = Path(str(filename)).name
    return filename if filename.lower().endswith(extension) else f"{filename}{extension}"


def create_csv(workspace, filename, rows):
    target = output_path(workspace, filename, ".csv")
    target.parent.mkdir(parents=True, exist_ok=True)
    headers, values = normalize_table_rows(rows)
    with target.open("w", encoding="utf-8", newline="") as output:
        writer = csv.writer(output)
        if headers:
            writer.writerow(headers)
        writer.writerows(values)
    relative = target.relative_to(Path(workspace).resolve())
    return f"Created CSV document: {relative}\nDownload: /agent_outputs/{target.name}"


def create_docx(workspace, filename, title, content):
    try:
        from docx import Document
    except ImportError:
        return "Cannot create DOCX: python-docx is not installed."

    target = output_path(workspace, filename, ".docx")
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    if title:
        document.add_heading(str(title), level=1)

    for block in str(content or "").splitlines():
        stripped = block.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith(("- ", "* ")):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(stripped)

    document.save(target)
    relative = target.relative_to(Path(workspace).resolve())
    return f"Created Word document: {relative}\nDownload: /agent_outputs/{target.name}"


def normalize_table_rows(rows):
    if not isinstance(rows, list):
        return [], []
    if not rows:
        return [], []
    if all(isinstance(row, dict) for row in rows):
        headers = []
        for row in rows:
            for key in row:
                if key not in headers:
                    headers.append(key)
        values = [[row.get(header, "") for header in headers] for row in rows]
        return headers, values
    values = [row if isinstance(row, list) else [row] for row in rows]
    return [], values


def create_xlsx(workspace, filename, rows, sheet_name="Sheet1"):
    try:
        from openpyxl import Workbook
    except ImportError:
        return "Cannot create XLSX: openpyxl is not installed."

    target = output_path(workspace, filename, ".xlsx")
    target.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = str(sheet_name or "Sheet1")[:31]
    headers, values = normalize_table_rows(rows)
    if headers:
        worksheet.append(headers)
    for row in values:
        worksheet.append(row)
    workbook.save(target)
    relative = target.relative_to(Path(workspace).resolve())
    return f"Created Excel workbook: {relative}\nDownload: /agent_outputs/{target.name}"


def run_command(workspace, command):
    try:
        arguments = shlex.split(command)
    except ValueError as error:
        return f"Invalid command: {error}"
    if not arguments:
        return "No command supplied."

    executable = Path(arguments[0]).name.lower()
    if executable in BLOCKED_COMMANDS:
        return f"Command blocked by policy: {executable}"
    if any(
        token in command
        for token in ("&&", "||", ";", "|", ">", "<", "`", "$(", "\n", "\r")
    ):
        return "Shell operators and redirections are not allowed."
    if executable in {"python", "python3", "node"} and any(
        flag in arguments for flag in ("-c", "-e", "--eval")
    ):
        return "Inline code execution is not allowed; write a workspace file first."

    try:
        completed = subprocess.run(
            arguments,
            cwd=workspace,
            capture_output=True,
            text=True,
            timeout=20,
            shell=False,
            env={
                "HOME": workspace,
                "PATH": os.environ.get("PATH", ""),
                "PYTHONDONTWRITEBYTECODE": "1",
            },
        )
    except FileNotFoundError:
        return f"Command not found: {arguments[0]}"
    except subprocess.TimeoutExpired:
        return "Command timed out after 20 seconds."

    output = "\n".join(
        part for part in (completed.stdout.strip(), completed.stderr.strip()) if part
    )
    return truncate(f"Exit code: {completed.returncode}\n{output or '(no output)'}")


def http_get(url, timeout=15):
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        raise ValueError("Only http and https URLs are allowed.")
    request = Request(url, headers={"User-Agent": "Mozilla/5.0 FlaskChatAgent/1.0"})
    with urlopen(request, timeout=timeout) as response:
        content_type = response.headers.get("Content-Type", "")
        body = response.read(MAX_FILE_BYTES).decode("utf-8", errors="replace")
    return content_type, body


def searxng_search(query, max_results):
    _, body = http_get(
        f"{SEARXNG_URL}/search?q={quote_plus(query)}&format=json"
    )
    payload = json.loads(body)
    results = []
    for result in payload.get("results", [])[:max_results]:
        results.append(
            {
                "title": result.get("title", ""),
                "url": result.get("url", ""),
                "snippet": result.get("content", ""),
                "source": "searxng",
            }
        )
    return results


def duckduckgo_search(query, max_results):
    _, body = http_get(f"https://html.duckduckgo.com/html/?q={quote_plus(query)}")
    parser = SearchResultParser()
    parser.feed(body)
    return [
        {**result, "source": "duckduckgo-fallback"}
        for result in parser.results[:max_results]
    ]


def web_search(query, max_results=5):
    max_results = max(1, min(int(max_results), 10))
    results = []
    searxng_error = None

    if SEARXNG_URL:
        try:
            results = searxng_search(query, max_results)
        except (HTTPError, URLError, TimeoutError, json.JSONDecodeError) as error:
            searxng_error = str(error)

    if not results:
        try:
            results = duckduckgo_search(query, max_results)
        except (HTTPError, URLError, TimeoutError) as error:
            if searxng_error:
                return (
                    f"SearXNG search failed: {searxng_error}\n"
                    f"Fallback search failed: {error}"
                )
            return f"Search failed: {error}"

    if not results:
        return "No search results found."
    return json.dumps(results, indent=2)


def fetch_url(url):
    content_type, body = http_get(url)
    if "html" in content_type:
        body = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", body)
        body = re.sub(r"(?s)<[^>]+>", " ", body)
        body = html.unescape(body)
        body = " ".join(body.split())
    return truncate(body)


TOOL_SCHEMAS = [
    {
        "type": "function",
        "function": {
            "name": "list_files",
            "description": "List files inside the private agent workspace.",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "default": "."}},
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": "Read a UTF-8 text file from the private agent workspace.",
            "parameters": {
                "type": "object",
                "required": ["path"],
                "properties": {"path": {"type": "string"}},
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_file",
            "description": "Create or replace a text file in the private agent workspace.",
            "parameters": {
                "type": "object",
                "required": ["path", "content"],
                "properties": {
                    "path": {"type": "string"},
                    "content": {"type": "string"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_command",
            "description": (
                "Run one non-interactive command in the private agent workspace. "
                "Shell operators, package installation, networking commands, git, and "
                "destructive commands are blocked."
            ),
            "parameters": {
                "type": "object",
                "required": ["command"],
                "properties": {"command": {"type": "string"}},
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_markdown",
            "description": "Create a Markdown document under outputs/ in the agent workspace.",
            "parameters": {
                "type": "object",
                "required": ["filename", "content"],
                "properties": {
                    "filename": {"type": "string"},
                    "content": {"type": "string"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_docx",
            "description": "Create a Word .docx document under outputs/ in the agent workspace.",
            "parameters": {
                "type": "object",
                "required": ["filename", "content"],
                "properties": {
                    "filename": {"type": "string"},
                    "title": {"type": "string"},
                    "content": {"type": "string"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_xlsx",
            "description": "Create an Excel .xlsx workbook under outputs/ in the agent workspace from rows.",
            "parameters": {
                "type": "object",
                "required": ["filename", "rows"],
                "properties": {
                    "filename": {"type": "string"},
                    "sheet_name": {"type": "string", "default": "Sheet1"},
                    "rows": {
                        "type": "array",
                        "items": {"type": "object"},
                    },
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_csv",
            "description": "Create a CSV file under outputs/ in the agent workspace from rows.",
            "parameters": {
                "type": "object",
                "required": ["filename", "rows"],
                "properties": {
                    "filename": {"type": "string"},
                    "rows": {
                        "type": "array",
                        "items": {"type": "object"},
                    },
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "Search the public internet and return result titles and URLs.",
            "parameters": {
                "type": "object",
                "required": ["query"],
                "properties": {
                    "query": {"type": "string"},
                    "max_results": {"type": "integer", "default": 5},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_url",
            "description": "Fetch and extract readable text from an HTTP or HTTPS URL.",
            "parameters": {
                "type": "object",
                "required": ["url"],
                "properties": {"url": {"type": "string"}},
            },
        },
    },
]


def execute_tool(name, arguments, workspace):
    tools = {
        "list_files": lambda: list_files(workspace, arguments.get("path", ".")),
        "read_file": lambda: read_file(workspace, arguments["path"]),
        "write_file": lambda: write_file(workspace, arguments["path"], arguments["content"]),
        "run_command": lambda: run_command(workspace, arguments["command"]),
        "create_markdown": lambda: create_markdown(
            workspace,
            argument_filename(arguments, "agent_output", ".md"),
            argument_text(arguments),
        ),
        "create_docx": lambda: create_docx(
            workspace,
            argument_filename(arguments, "agent_output", ".docx"),
            arguments.get("title", ""),
            argument_text(arguments),
        ),
        "create_xlsx": lambda: create_xlsx(
            workspace,
            argument_filename(arguments, "agent_output", ".xlsx"),
            arguments["rows"],
            arguments.get("sheet_name", "Sheet1"),
        ),
        "create_csv": lambda: create_csv(
            workspace,
            argument_filename(arguments, "agent_output", ".csv"),
            arguments["rows"],
        ),
        "web_search": lambda: web_search(
            arguments["query"], arguments.get("max_results", 5)
        ),
        "fetch_url": lambda: fetch_url(arguments["url"]),
    }
    if name not in tools:
        return f"Unknown tool: {name}"
    try:
        return tools[name]()
    except (KeyError, TypeError, ValueError) as error:
        return f"Invalid tool arguments: {error}"
    except (HTTPError, URLError, TimeoutError) as error:
        return f"Network tool failed: {error}"


def json_objects_from_text(content):
    objects = []
    start = None
    depth = 0
    in_string = False
    escape = False

    for index, character in enumerate(content or ""):
        if in_string:
            if escape:
                escape = False
            elif character == "\\":
                escape = True
            elif character == '"':
                in_string = False
            continue

        if character == '"':
            in_string = True
        elif character == "{":
            if depth == 0:
                start = index
            depth += 1
        elif character == "}" and depth:
            depth -= 1
            if depth == 0 and start is not None:
                objects.append(content[start : index + 1])
                start = None

    return objects


def tool_calls_from_content(content):
    content = (content or "").strip()
    if not content:
        return []

    candidates = []
    try:
        payload = json.loads(content)
        candidates.extend(payload if isinstance(payload, list) else [payload])
    except json.JSONDecodeError:
        pass

    for object_text in json_objects_from_text(content):
        try:
            candidates.append(json.loads(object_text))
        except json.JSONDecodeError:
            continue

    tool_calls = []
    for candidate in candidates:
        if not isinstance(candidate, dict):
            continue

        name = (
            candidate.get("name")
            or candidate.get("tool")
            or candidate.get("tool_name")
            or candidate.get("function")
        )
        arguments = (
            candidate.get("arguments")
            or candidate.get("parameters")
            or candidate.get("args")
            or {}
        )

        if isinstance(name, dict):
            arguments = name.get("arguments") or arguments
            name = name.get("name")
        if not isinstance(name, str):
            continue
        if isinstance(arguments, str):
            try:
                arguments = json.loads(arguments)
            except json.JSONDecodeError:
                arguments = {}
        if not isinstance(arguments, dict):
            arguments = {}

        tool_calls.append({"function": {"name": name, "arguments": arguments}})
    return tool_calls


def looks_like_tool_planning(content):
    text = (content or "").lower()
    if not text:
        return False
    tool_names = [
        schema["function"]["name"].lower()
        for schema in TOOL_SCHEMAS
        if schema.get("function", {}).get("name")
    ]
    return (
        any(tool_name in text for tool_name in tool_names)
        and any(marker in text for marker in ("parameters", "arguments", "i will use", "tool"))
    )


def load_agent_instructions(
    base_instructions_file,
    orchestrator_instructions_file,
    skills_dir,
    workspace,
):
    orchestrator = Path(orchestrator_instructions_file).read_text(encoding="utf-8")
    base = Path(base_instructions_file).read_text(encoding="utf-8")
    skill_sections = []
    for skill_file in sorted(Path(skills_dir).glob("*.md")):
        skill_sections.append(
            f"\n## Skill: {skill_file.stem}\n{skill_file.read_text(encoding='utf-8')}"
        )
    return (
        f"{orchestrator.strip()}\n\n"
        f"{base.strip()}\n\n"
        f"Your private workspace is {Path(workspace).resolve()}. "
        "All file and command tools are confined there.\n"
        + "\n".join(skill_sections)
    )


def run_agent(
    model,
    messages,
    ollama_chat,
    workspace,
    base_instructions_file,
    orchestrator_instructions_file,
    skills_dir,
    max_steps=8,
    status_callback=None,
):
    agent_messages = [
        {
            "role": "system",
            "content": load_agent_instructions(
                base_instructions_file,
                orchestrator_instructions_file,
                skills_dir,
                workspace,
            ),
        },
        *messages,
    ]
    trace = []

    for step in range(1, max_steps + 1):
        if status_callback:
            status_callback(
                {
                    "state": "thinking",
                    "step": step,
                    "message": (
                        f"Orchestrator is deciding whether step {step} needs "
                        "a direct answer or tools."
                    ),
                }
            )
        response_message = ollama_chat(model, agent_messages, TOOL_SCHEMAS)
        agent_messages.append(response_message)
        tool_calls = response_message.get("tool_calls") or tool_calls_from_content(
            response_message.get("content")
        )
        if not tool_calls:
            content = (response_message.get("content") or "").strip()
            if looks_like_tool_planning(content):
                agent_messages.append(
                    {
                        "role": "user",
                        "content": (
                            "That response looked like internal tool planning rather "
                            "than a user-facing answer. If a tool is needed, call it "
                            "through the tool interface with valid arguments. If the "
                            "task is complete, write the final answer for the main "
                            "chat in clear Markdown. Do not print JSON tool-call "
                            "objects or orchestration notes."
                        ),
                    }
                )
                if status_callback:
                    status_callback(
                        {
                            "state": "thinking",
                            "step": step,
                            "message": (
                                "Agent output looked like tool planning; asking the "
                                "orchestrator to produce a proper chat response."
                            ),
                        }
                    )
                continue
            if status_callback:
                status_callback(
                    {
                        "state": "complete",
                        "step": step,
                        "message": "Agent finished without calling another tool.",
                    }
                )
            return content or "The agent completed without a text response.", trace

        for tool_call in tool_calls:
            function = tool_call.get("function", {})
            name = function.get("name", "")
            arguments = function.get("arguments") or {}
            if isinstance(arguments, str):
                try:
                    arguments = json.loads(arguments)
                except json.JSONDecodeError:
                    arguments = {}
            if status_callback:
                status_callback(
                    {
                        "state": "running_tool",
                        "step": step,
                        "tool": name,
                        "arguments": arguments,
                        "message": describe_tool_start(name, arguments),
                    }
                )
            result = execute_tool(name, arguments, workspace)
            trace_entry = {
                "step": step,
                "tool": name,
                "arguments": arguments,
                "status": describe_tool_result(name, result),
                "result": truncate(result, 1_000),
            }
            trace.append(trace_entry)
            if status_callback:
                status_callback(
                    {
                        "state": "tool_complete",
                        "step": step,
                        "tool": name,
                        "arguments": arguments,
                        "message": trace_entry["status"],
                    }
                )
            agent_messages.append(
                {"role": "tool", "tool_name": name, "content": str(result)}
            )

        agent_messages.append(
            {
                "role": "user",
                "content": (
                    "Tool activity above is private work. If you have enough "
                    "information, hand off to the output writer now: write the "
                    "final user-facing answer for the main chat in clear Markdown. "
                    "Do not expose raw tool-call JSON, parameters, or internal "
                    "orchestration notes. If another tool is genuinely required, "
                    "call it through the tool interface."
                ),
            }
        )

    agent_messages.append(
        {
            "role": "user",
            "content": (
                "You have reached the tool-step limit. Stop using tools and give the "
                "best final answer possible from the work completed."
            ),
        }
    )
    final_message = ollama_chat(model, agent_messages, [])
    if status_callback:
        status_callback(
            {
                "state": "complete",
                "step": max_steps,
                "message": "Agent stopped at the step limit and wrote a final answer.",
            }
        )
    return (
        (final_message.get("content") or "Agent stopped at its step limit.").strip(),
        trace,
    )


def describe_tool_start(name, arguments):
    if name == "web_search":
        query = arguments.get("query", "")
        return f'Searching the web for "{query}" via SearXNG.'
    if name == "fetch_url":
        return f"Fetching URL: {arguments.get('url', '')}"
    if name == "run_command":
        return f"Running command: {arguments.get('command', '')}"
    if name == "read_file":
        return f"Reading file: {arguments.get('path', '')}"
    if name == "write_file":
        return f"Writing file: {arguments.get('path', '')}"
    if name == "create_markdown":
        return f"Creating Markdown document: {arguments.get('filename', '')}"
    if name == "create_docx":
        return f"Creating Word document: {arguments.get('filename', '')}"
    if name == "create_xlsx":
        return f"Creating Excel workbook: {arguments.get('filename', '')}"
    if name == "create_csv":
        return f"Creating CSV document: {arguments.get('filename', '')}"
    if name == "list_files":
        return f"Listing files: {arguments.get('path', '.')}"
    return f"Running tool: {name}"


def describe_tool_result(name, result):
    result_text = str(result)
    if name == "web_search":
        if result_text.startswith("SearXNG search failed"):
            return "Web search failed on SearXNG and the fallback."
        if result_text.startswith("Search failed"):
            return "Web search failed."
        if result_text == "No search results found.":
            return "Web search completed with no results."
        try:
            parsed = json.loads(result_text)
        except json.JSONDecodeError:
            return "Web search completed."
        source = parsed[0].get("source") if parsed else None
        if source == "searxng":
            return f"Web search returned {len(parsed)} result(s) from SearXNG."
        if source == "duckduckgo-fallback":
            return (
                f"Web search returned {len(parsed)} result(s) from the fallback search."
            )
        return f"Web search returned {len(parsed)} result(s)."
    if name == "fetch_url":
        return f"Fetched {len(result_text)} character(s)."
    if name == "run_command":
        first_line = result_text.splitlines()[0] if result_text else "Command finished."
        return first_line
    return result_text.splitlines()[0] if result_text else f"{name} completed."
